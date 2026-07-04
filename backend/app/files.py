"""Real document parsing for manual uploads: PDF, DOCX, XLSX, CSV, TXT, MD."""
from __future__ import annotations

import csv
import io


def parse_file(filename: str, data: bytes) -> str:
    name = filename.lower()
    try:
        if name.endswith(".pdf"):
            return _pdf(data)
        if name.endswith(".docx"):
            return _docx(data)
        if name.endswith((".xlsx", ".xlsm")):
            return _xlsx(data)
        if name.endswith(".csv"):
            return _csv(data)
    except Exception:
        # Fall through to best-effort text decode.
        pass
    return data.decode("utf-8", errors="ignore")


def _pdf(data: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _docx(data: bytes) -> str:
    import docx
    doc = docx.Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _xlsx(data: bytes) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    lines = []
    for ws in wb.worksheets:
        lines.append(f"# Sheet: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                lines.append(", ".join(cells))
    return "\n".join(lines)


def _csv(data: bytes) -> str:
    text = data.decode("utf-8", errors="ignore")
    reader = csv.reader(io.StringIO(text))
    return "\n".join(", ".join(row) for row in reader)
